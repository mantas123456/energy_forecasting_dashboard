# phase_2_modeling_pipeline/scripts/09_generate_phase2_report.py

"""
09_generate_phase2_report.py
-----------------------------
Generates a metric-system-compliant .docx report for Phase 2 evaluation,
now including feature descriptions.

Author: Mantas Valantinavicius
"""

from docx import Document
from docx.shared import Inches
from pathlib import Path
import pandas as pd
import logging

logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")

RESULTS_DIR = Path("phase_2_modeling_pipeline/results")
PLOTS_DIR = RESULTS_DIR / "plots"
REPORT_PATH = RESULTS_DIR / "Phase2_Model_Evaluation_Report.docx"

FEATURES = [
    ("hour", "Hour of the day (0–23)", "Integer", "Captures daily consumption patterns"),
    ("day_of_week", "Day of week (0 = Monday, 6 = Sunday)", "Integer", "Captures weekly usage trends"),
    ("month", "Month of year (1–12)", "Integer", "Captures seasonal effects"),
    ("is_weekend", "1 = weekend, 0 = weekday", "Binary", "Distinguishes weekend usage behavior"),
    ("hour_sin", "Sine transformation of hour", "Float", "Captures cyclical daily pattern"),
    ("hour_cos", "Cosine transformation of hour", "Float", "Complements `hour_sin` for full cycle encoding"),
    ("dow_sin", "Sine transformation of day_of_week", "Float", "Captures weekly pattern phase"),
    ("dow_cos", "Cosine transformation of day_of_week", "Float", "Complements `dow_sin`"),
    ("temperature_C", "Ambient temperature", "Degrees Celsius (°C)", "Relates to heating/cooling needs"),
    ("lag_1h", "Energy use 1 hour ago", "kWh", "Captures short-term trend"),
    ("lag_24h", "Energy use same hour previous day", "kWh", "Captures daily recurrence"),
    ("roll_mean_24h", "Rolling 24h mean energy use", "kWh", "Smooths noise and shows recent trends"),
]

def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)

def add_paragraph(doc, text):
    doc.add_paragraph(text)

def add_image(doc, image_path, width_inches=6, caption=None):
    if image_path.exists():
        doc.add_picture(str(image_path), width=Inches(width_inches))
        doc.add_paragraph(f"Figure: {caption or image_path.name}")
    else:
        doc.add_paragraph(f"⚠ Image not found: {image_path}")

def add_evaluation_table(doc, csv_path):
    if not csv_path.exists():
        doc.add_paragraph("⚠ Evaluation summary not found.")
        return

    df = pd.read_csv(csv_path)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Light Grid'

    # Header
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = col

    # Rows
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = f"{value:.4f}" if isinstance(value, float) else str(value)

    doc.add_paragraph("Note: All error metrics (RMSE, MAE, MAPE) are reported in kilowatt-hours (kWh).")

def add_feature_description_table(doc):
    add_heading(doc, "2.1 Feature Descriptions", level=2)
    add_paragraph(doc, "The following table summarizes all engineered input features used in model training:")

    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Feature Name"
    hdr_cells[1].text = "Description"
    hdr_cells[2].text = "Unit / Type"
    hdr_cells[3].text = "Purpose in Model"

    for feature, desc, unit, purpose in FEATURES:
        row = table.add_row().cells
        row[0].text = feature
        row[1].text = desc
        row[2].text = unit
        row[3].text = purpose

def generate_report():
    doc = Document()

    # Title
    doc.add_heading("Phase 2: Model Evaluation Report", level=0)
    doc.add_paragraph("Author: Mantas Valantinavicius")
    doc.add_paragraph("Date: Auto-generated")

    # Overview
    add_heading(doc, "1. Project Overview")
    add_paragraph(doc, (
        "This report summarizes the results of Phase 2 of the energy forecasting dashboard project. "
        "The objective was to train and compare multiple time-series forecasting models to predict hourly energy consumption. "
        "All modeling is performed using metric system inputs and outputs: energy in kilowatt-hours (kWh), "
        "temperature in degrees Celsius (°C), and time in hourly intervals."
    ))

    # Methodology
    add_heading(doc, "2. Methodology")
    add_paragraph(doc, (
        "The synthetic dataset was enriched with temporal, cyclical, and lag-based features. "
        "Models were trained on 80% of the dataset and evaluated using RMSE, MAE, and MAPE, all expressed in kilowatt-hours (kWh). "
        "XGBoost was configured with 100 estimators and a maximum depth of 5. Prophet included daily and weekly seasonality. "
        "Linear Regression was used as a baseline model for comparison."
    ))

    # NEW SECTION: Feature Descriptions
    add_feature_description_table(doc)

    # Evaluation Table
    add_heading(doc, "3. Model Performance Summary")
    add_evaluation_table(doc, RESULTS_DIR / "model_evaluation_summary.csv")

    # Prediction Visuals
    add_heading(doc, "4. Actual vs Predicted Plots")
    add_paragraph(doc, "Note: All vertical axes represent energy use in kilowatt-hours (kWh).")
    for model in ["prophet", "xgboost", "linear"]:
        add_image(doc, PLOTS_DIR / f"{model}_prediction_plot.png", caption=f"{model.title()} model predictions")

    # XGBoost Feature Importance
    add_heading(doc, "5. Feature Importance (XGBoost)")
    add_image(doc, PLOTS_DIR / "xgboost_feature_importance.png", caption="Relative importance of each feature")

    # Linear Regression Coefficients
    add_heading(doc, "6. Feature Coefficients (Linear Regression)")
    add_image(doc, PLOTS_DIR / "linear_feature_coefficients.png", caption="Signed influence of each input feature")

    # Prophet Internals
    add_heading(doc, "7. Prophet Model Components")
    add_image(doc, PLOTS_DIR / "prophet_components_plot.png", caption="Prophet trend and daily seasonality components")
    add_paragraph(doc, (
        "The above plot illustrates the components learned by Prophet, including overall trend and repeating daily cycles. "
        "These cycles reflect typical human activity patterns, such as energy use peaks during working hours."
    ))

    # Conclusion
    add_heading(doc, "8. Conclusion")
    add_paragraph(doc, (
        "Among the evaluated models, XGBoost produced the best performance, achieving the lowest error across all metrics. "
        "Its effectiveness is attributed to strong use of time-based features and lagged values. "
        "Prophet offered strong interpretability through trend and seasonality decomposition. "
        "Linear Regression served as a useful but limited benchmark. "
        "All models were trained and assessed using metric units for real-world applicability."
    ))

    # Save
    doc.save(REPORT_PATH)
    logging.info(f"📄 Metric system report with feature descriptions saved: {REPORT_PATH.resolve()}")


if __name__ == "__main__":
    generate_report()
